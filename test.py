
context = {"first_name": "Benjamin",
           "last_name":"Sliosberg",
           "job_position":"Data scientist",
           "years_experience":"3",
           "professional_summary":"""Data science4 machin machine 
           truc 
           hahaha""",
           "languages": 
            [  
                {"name": "English", "level": "Native"},  
                {"name": "French", "level": "Intermediate"},  
                {"name": "Spanish", "level": "Beginner"}  
            ],
            "experiences" : 
            [
                {
                    "company_name":"EFOR",
                    "position":"data scientist",
                    "start_month":"01",
                    "start_year":"2020",
                    "end_month":"10",
                    "end_year":"2021",
                    "responsibilities": 
                        [
                            'Washed dogs',
                            'kapuera Master'
                        ]
                },
                {
                    "company_name":"SOLADIS",
                    "position":"Data Analyst",
                    "start_month":"01",
                    "start_year":"2020",
                    "end_month":"10",
                    "end_year":"2021",
                    "responsibilities": 
                        [
                            'deepdived in the documentation',
                            'Washed cats',
                            'kapuera begginer'
                        ]
                }
            ],
            "skills":
            [
                {"name":"BEING FUN"},
                {"name":"LOL"}
            ],
            "softwares":
            [
                {"name":"Python"},
                {"name":"R"},
            ],
            "standards":
            [
                {"name":"STANDARDS ? ME //"},
                {"name":"ISO 6623"}
            ],

            "educations":
                [
                
                    {
                        "name":"Master Computer Science (M. CS.)",
                        'field_of_study':'Data science',
                        'university_name':"ESILV",
                        'graduation_year': "2023"
                    },
                    {
                        "name":"Master Computer Science (M. CS.)",
                        'field_of_study':'Data science 2',
                        'university_name':"CSULB",
                        'graduation_year': "2025"
                    }, 
                ],
            "certifications":
                [
                    {"name":"ASSR2"},
                    {"name":"DataIKU DEVOPS"},
                ]

                
            }
context = {
"first_name": "Benjamin",
"last_name": "Sliosberg",
"job_position": "Consultant Data Scientist / Developer",
"years_experience": "3",
"professional_summary": "Bilingual engineer with 3 years of experience in predictive models and software development seeking a Machine Learning / BackEnd Developer short mission. Internationally minded, with a strong inclination towards geographical mobility. Available from January, 2025.",
"languages": [
{ "name": "English", "level": "Native" },
{ "name": "French", "level": "Native" }
],
"experiences": [
{
"company_name": "Efor Group",
"position": "Consultant Data Scientist - Developer",
"start_month": "09",
"start_year": "2023",
"end_month": "",
"end_year": "",
"responsibilities": [
"Built and deployed a full-stack application on Azure using a Django and PostgreSQL backend in 3 months (Nestlé).",
"Designed an Azure Functions–based code generator enabling production staff to independently generate 50 codes; improved productivity by ~95% (Nestlé).",
"Developed and maintained CI/CD Azure Pipelines for 3 Docker apps and Azure Functions (Nestlé).",
"Delivered data visualizations and report-building tools to onsite and statistical engineers, reducing report generation time by 5× (Nestlé).",
"Built a Streamlit stability-study platform with statistical models, cutting report generation time from 7 days to 45 minutes and expediting FDA submissions by ~80% (Biomérieux).",
"Automated a manual statistical analysis process via a custom application, boosting efficiency by ~85% (Biomérieux).",
"Identified a 157-variant genetic signature predictive of sudden death in newborns through genomic data analysis (Lyon CHU).",
"Devised a method to predict cetacean occurrences using NASA satellite data and OBIS sightings; created a Power BI dashboard integrating Copernicus and OBIS data (Copernicus Marine)."
]
},
{
"company_name": "LAMM - Iconik (Emotional VR)",
"position": "Machine Learning Engineer",
"start_month": "01",
"start_year": "2023",
"end_month": "07",
"end_year": "2023",
"responsibilities": [
"Researched a neural network to predict real-time human emotional states from biometric data.",
"Developed a state-of-the-art peak detection approach for biometric signals using CNN methodology."
]
},
{
"company_name": "Électricité de France (EDF) R&D Lab",
"position": "Software Engineer",
"start_month": "03",
"start_year": "2022",
"end_month": "08",
"end_year": "2022",
"responsibilities": [
"Developed multi-agent systems modeling French energy consumption to inform strategies during the 2022 winter crisis.",
"Deployed a web application for energy models with three interactive simulations using agent behaviors based on survey and bias-study data."
]
}
],
"skills": [
{ "name": "Machine Learning" },
{ "name": "Predictive Modeling" },
{ "name": "Data Visualization" },
{ "name": "Web Application Development" },
{ "name": "DevOps & CI/CD" },
{ "name": "Statistical Analysis" },
{ "name": "Genomic Data Analysis" },
{ "name": "Time-Series & Peak Detection" },
{ "name": "Multi-Agent Modeling" },
{ "name": "Dashboarding" },
{ "name": "Prompt Engineering" },
{ "name": "Communication" },
{ "name": "Group Dynamics" },
{ "name": "Team Motivation" }
],
"softwares": [
{ "name": "Python" },
{ "name": "R" },
{ "name": "MATLAB" },
{ "name": "Scikit-learn" },
{ "name": "TensorFlow" },
{ "name": "Keras" },
{ "name": "Power BI" },
{ "name": "Dataiku" },
{ "name": "Streamlit" },
{ "name": "Django" },
{ "name": "Azure DevOps" },
{ "name": "Docker" },
{ "name": "MySQL" },
{ "name": "PostgreSQL" },
{ "name": "LaTeX" }
],
"standards": [],
"educations": [
{
"name": "Masters Specialization - ML & Data Visualization",
"field_of_study": "Machine Learning & Data Visualization",
"university_name": "CSULB (California State University, Long Beach)",
"graduation_year": ""
},
{
"name": "Engineering Masters",
"field_of_study": "Engineering (Research minor)",
"university_name": "ESILV (École Supérieure d'Ingénieurs Léonard-de-Vinci)",
"graduation_year": ""
}
],
"certifications": [
{ "name": "Dataiku MLOps Practitioner" },
{ "name": "Dataiku Advanced Developer" }
]
}

# from docxtpl import Listing  
# from docxtpl import RichText  
# from docxtpl import DocxTemplate, RichTextParagraph  

  
# for experience in context["experiences"]:  
#     bullet_text = "\n".join([f"• {resp}" for resp in experience["responsibilities"]])  
#     experience["bullet_list"] = Listing(bullet_text)


# # rt = RichText()  
# # for i, education in enumerate(context["educations"]):  
# #     print(i,education)
# #     if i > 0:  
# #         rt.add("\a")  # Add paragraph break between items  
# #     rt.add("• ")  # Add bullet character  
# #     rt.add(education['name'], bold=True)  # Add responsibility text in bold  
# #     rt.add(" - ")
# #     rt.add(education['field_of_study'])
# #     rt.add(", ")
# #     rt.add(education['university_name'])
# #     rt.add(education['graduation_year'])
# # education['bullet_list'] = rt

# rtp = RichTextParagraph()  
# for education in context['educations']:  
#     education_text = f"{education['name']} - {education['field_of_study']}, {education['university_name']} ({education['graduation_year']})"  
#     rtp.add(education_text, parastyle="SquareBullet")

# tpl.render(context)

# tpl.save("Efor template filled.docx")


from pathlib import Path
import json

from docxtpl import DocxTemplate, Listing
from docx.enum.text import WD_ALIGN_PARAGRAPH


context = json.loads(Path("output_resume.json").read_text(encoding="utf-8"))

doctemplate = r"templates/EFOR TEMPLATE.docx"

tpl = DocxTemplate(doctemplate)


for experience in context["experiences"]:
    bullet_text = "\n".join(f"• {resp}" for resp in experience["responsibilities"])
    experience["bullet_list"] = Listing(bullet_text)

edu_sd = tpl.new_subdoc()
for edu in context["educations"]:
    # Use a bullet style that exists in the template.
    p = edu_sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(f"• {edu["name"]}")
    r.bold = True
    p.add_run(f" – {edu['field_of_study']}, {edu['university_name']} ({edu['graduation_year']})")

# Put it where your template expects it
context["education"] = {"bullet_list": edu_sd}

sd = tpl.new_subdoc()
for c in context["certifications"]:
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("• ")
    p.add_run(c["name"])

context["certifications"] = {"bullet_list": sd}

# 4) render and save
tpl.render(context)
tpl.save("Efor template filled.docx")