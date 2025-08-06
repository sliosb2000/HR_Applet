import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from datetime import datetime
import re 

def extract_placeholders(doc):
    """
    Scan the entire document for tags surrounded by {{ and }},
    returning a sorted list of unique placeholder keys.
    """
    tags = set()
    # Check paragraphs
    for para in doc.paragraphs:
        found = re.findall(r"\{\{(.*?)\}\}", para.text)
        tags.update(found)
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found = re.findall(r"\{\{(.*?)\}\}", cell.text)
                tags.update(found)
    return sorted(tags)


# --- Replacement utility ---
def replace_placeholders(doc, mapping):
    """
    Replace {{KEY}} tags in paragraphs and table cells by setting element.text directly.
    This approach consolidates all runs, avoiding per-run iteration.
    """
    for key, val in mapping.items():
        if val not in [""," ","  "]:
            placeholder = f"{{{{{key}}}}}"
            # Replace in paragraphs
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, val)
                    # Ensure runs adopt Calibri
                    for run in para.runs:
                        run.font.name = "Calibri"
                        # run.font.size = 12


# Define a styling function: highlights Key & Description red if Value is blank or equals Description
def highlight_keys_desc(row):
    v = str(row["Value"])
    desc = str(row["Description"])
    # Condition: empty/whitespace OR identical to Description
    cond = (v.strip() == "") or (v == " ") or (v == desc)
    # Apply red background to Key & Description, else no style
    return ["background-color: #E57373" if col in ["Key", "Description"] and cond else "" for col in row.index]

def highlight_placeholders_background(doc, fill_color="FF0000"):
    """
    For every run whose text matches {{...}}, add a w:shd element
    to its rPr, giving it a background of `fill_color` (hex, no #).
    """
    pattern = re.compile(r"\{\{.*?\}\}")

    def shade_run(run):
        rPr = run._element.get_or_add_rPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        rPr.append(shd)

    for para in doc.paragraphs:
        if pattern.search(para.text):
            for run in para.runs:
                shade_run(run)



# --- Page Config ---
st.set_page_config(page_title="HR Automation Tool", layout="wide")

st.markdown("HR Automation Tool")
st.sidebar.markdown("Page 1: Data Input & Editing")

# --- 1. File upload: Source data (CSV, XLSX, or Word table) ---
data_file = st.file_uploader(
    label="Upload source data (CSV, XLSX, or Word .docx)",
    type=["csv", "xlsx", "docx"],
    help="CSV/XLSX should have columns [Key, Value]. Word should contain a single 2-col table with labels in col1 and blanks in col2.",
    key="data_upload"
)


#IDs of Cells of interest in the Offer_letter_input file:
cells_of_interest = {
    "full_name":1,
    "first_name":2,
    "last_name":3,
    "current_address":5,
    "phone_number":6,
    "email":7,
    "job_title":20,
    "start_date":21,
    "supervisor_full_name":24,
    "main_duties":29,
    "annual_gross_salary":22,
    #contingent info
    "client_name":31,
    "client_location":32,
    "client_response_date":33,
    "deadline_response_date":34    
}

# Initialize empty dataframe
df_information = [
    {"Key": "full_name",              "Description": "Full legal name (as it appears on passport or official documents)",                                 "Value": ""},
    {"Key": "first_name",             "Description": "First name (if ignored will use the first part of full name)",                                        "Value": ""},
    {"Key": "last_name",              "Description": "Last name (if ignored will use the last part of full name)",                                             "Value": ""},
    {"Key": "current_address",        "Description": "Current address",                                                                                    "Value": ""},
    {"Key": "phone_number",           "Description": "Phone number",                                                                                     "Value": ""},
    {"Key": "email",                  "Description": "Email address",                                                                                      "Value": ""},
    {"Key": "job_title",              "Description": "Job title",                                                                                           "Value": ""},
    {"Key": "start_date",             "Description": "Start date",                                                                                       "Value": ""},
    {"Key": "supervisor_full_name",   "Description": "Manager",                                                                                           "Value": ""},
    {"Key": "main_duties",            "Description": "Example: Support Commissioning, Qualification, and Validation (CQV) activities related to vaccine manufacturing processes, including Cleaning Validation, Equipment Qualification, and Process Validation. Author and execute protocols and reports in compliance with regulatory standards.",  "Value": ""},
    {"Key": "annual_gross_salary",    "Description": "Salary proposed (per hours for contractors or annual gross salary for W2)",                      "Value": ""},
    #contingent info
    {"Key": "client_name",            "Description": "Client Name",                                                                                        "Value": ""},
    {"Key": "client_location",        "Description": "Client Location",                                                                                    "Value": ""},
    {"Key": "client_response_date",   "Description": "Client Response Date",                                                                               "Value": ""},
    {"Key": "deadline_response_date", "Description": "Deadline Response Date",                                                                              "Value": ""},
]

# Build the DataFrame
if "df" not in st.session_state:
    st.session_state.df = pd.DataFrame(df_information)
st.session_state.df.index = st.session_state.df.Key


if data_file is not None:
    fname = data_file.name.lower()
    if fname.endswith(".xlsx"):
        st.session_state.df = pd.read_excel(data_file)
    elif fname.endswith(".csv"):
        # Try comma then semicolon separators
        try:
            st.session_state.df = pd.read_csv(data_file, sep=",")
            if st.session_state.df.shape[1] < 2:
                raise ValueError("Only one column detected, retrying with semicolon separator.")
        except Exception:
            st.session_state.df = pd.read_csv(data_file, sep=";")
    else:
        # Parse Word table
        doc = Document(data_file)
        tables = doc.tables
        # print(f"Found {len(tables)} table(s) in {data_file!r}")

        pairs = []
        for ti, table in enumerate(tables, start=1):
            # print(f" – Table {ti}: {len(table.rows)} rows")
            for ri, row in enumerate(table.rows, start=1):
                # Skip any rows that don’t have at least 2 cells
                if len(row.cells) < 2:
                    continue
                key   = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                # Only append if there’s some content
                pairs.append({'Description': key, 'Value': value})

        df_raw = pd.DataFrame(pairs)
        
        # Filter df_raw using cells_of_interest dictionary
        # Create a new dataframe with only the rows specified in cells_of_interest
        filtered_rows = []
        if len(df_raw) < max(cells_of_interest.values()):
            st.error("Wrong Word template used. Please use the correct template or fill out table manually.")  
            filtered_rows = df_information
            st.session_state.df = pd.DataFrame(filtered_rows)
        else:
            for key, row_index in cells_of_interest.items():
                if row_index < len(df_raw):
                    row_data = df_raw.iloc[row_index].copy()
                    row_data.name = key  # Name the row with the dictionary key
                    filtered_rows.append(row_data)
        
            # Create the filtered dataframe
            df_filtered = pd.DataFrame(filtered_rows)
            
            # Use the filtered dataframe for editing
            st.session_state.df = df_filtered
            df_filtered["Key"] = df_filtered.index
            st.session_state.df = df_filtered[["Key","Description","Value"]] #reorder columns
        st.session_state.df = st.session_state.df.astype({'Key': 'string', 'Value': 'string', 'Description': 'string'})
        st.session_state.df.index = st.session_state.df.Key



# --- 2. Editable table for user corrections ---
if not st.session_state.df.empty:
    st.subheader("Review & Edit Data")
    st.button("refresh table",icon="🔄")

    df = st.session_state.df

    df.loc["first_name", "Value"] = (
        df.loc["first_name", "Value"].strip()
        if df.loc["first_name", "Value"].strip() != ""
        else df.loc["full_name", "Value"].split(" ")[0]
    )

    df.loc["last_name", "Value"] = (
        df.loc["last_name", "Value"].strip()
        if df.loc["last_name", "Value"].strip() != ""
        else df.loc["full_name", "Value"].split(" ")[-1]
    )
    
    styled_df = df.style.apply(highlight_keys_desc, axis=1)
    edited_df = st.data_editor(styled_df, 
                               use_container_width=True,
                               column_config={
                                   "Key": st.column_config.TextColumn(
                                       width="small",
                                       help="Clef de la cellule d'intérêt"
                                   ),
                                   "Description": st.column_config.TextColumn(
                                       width="small",
                                       help="Description pour l'utilisateur"
                                   ),
                                   "Value": st.column_config.TextColumn(
                                       width="large",
                                       help="Valeur de la cellule d'intérêt"
                                   ),
                               },
                               disabled=["Key","Description"],
                               hide_index=True)
    
    st.session_state.df = edited_df

else:
    st.info("Upload a CSV, XLSX, or DOCX file to begin.")
    st.stop()

type_document = st.selectbox("Type de document",["Employment Contract","Contingent Offer","Custom Input Document"])

if type_document == "Employment Contract":
    # st.write("Employment Contract")
    template_doc = Document("data/Template_Employment_Contract_EFOR.docx")
elif type_document == "Contingent Offer":
    # st.write("Contingent Offer")
    template_doc = Document("data/Template_Contingent_Offer_EFOR.docx")
elif type_document == "Custom Input Document":
    st.write("Please Input a Custom Template Document")

    # --- 3. Template upload or default ---
    template_doc = st.file_uploader(
        label="Upload Word template (.docx)", 
        type=["docx"],
        key="template_upload"
    )

        


# --- 4. Generate filled document ---
if st.button("Generate & Download Word Document"):
    if template_doc is None:
        st.error("Please upload a template document.")
        st.stop()
    try:

        
        # Build mapping dict from edited_df
        mapping = {row.Key: str(row.Value) for row in edited_df.itertuples(index=False)}
        # Add current_date if needed
        mapping["current_date"] = datetime.today().strftime("%d/%m/%Y")


        # Perform placeholder replacement
        replace_placeholders(template_doc, mapping)

        ununsed_Keys = extract_placeholders(template_doc)
        highlight_placeholders_background(template_doc)
        if len(ununsed_Keys) > 0:
            st.info(f"The Document is Generated but missing the following information, please input manually in the file : \n\n{"\n\n".join(ununsed_Keys)}")

        # Save to buffer and offer download
        buffer = BytesIO()
        template_doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="Download Filled Document",
            data=buffer,
            file_name=f"{"_".join(type_document.split(" "))}_{"_".join(edited_df["Value"]["full_name"].split(" "))}_{"_".join(edited_df["Value"]["job_title"].split(" "))}.docx",

            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Error generating document: {e}")



